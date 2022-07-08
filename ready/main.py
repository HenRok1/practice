import openpyxl
from docx import Document
import os
import shutil
import csv
import pandas as pd

def flag_check(flag_csv):
    if flag_csv == 0:
        practitioner_data = [
                ("id", "Документ", "Основание", "Файл №1", "Файл №2", "Файл №3", "Файл №4")
            ]
        with open("practitioner.csv", "a", newline='', encoding='cp1251', errors="ignore") as csvfile:
            practitioner = csv.writer(csvfile, delimiter=";")
            practitioner.writerows(
                practitioner_data
            )

def path_check(num_of_docs):
    for num_of_docs in range(0, num_of_docs):

        folder_save = input("Введите название папки для сохранения:\n")
        os.mkdir(folder_save)

        file_name1 = input("Введите имя файла №1, в котором содержится ПУСТОЙ СКРИН ИЗ ГОСТА (с расширением .pdf или .jpg):\n")
        if file_name1 != "":
            path_file1 = os.path.abspath(file_name1)
            Doc1 = shutil.copy(file_name1, folder_save)

        else:
            path_file1 = " "
        file_name2 = input("Введите имя файла №2, в котором содержится ПУСТОЙ ШАБЛОН (с расширением .doc или .docx):\n")
        if file_name2 != "":
            path_file2 = os.path.abspath(file_name2)
            Doc2 = shutil.copy(file_name2, folder_save)

        else:
            path_file2 = " "
        file_name3 = input("Введите имя файла №3, в котором содержится ЗАПОЛНЕННЫЙ ДОКУМЕНТ (с расширением  .pdf или .jpg):\n")
        if file_name3 != "":
            path_file3 = os.path.abspath(file_name3)
            Doc3 = shutil.copy(file_name3, folder_save)

        else:
            path_file3 = " "
        file_name4 = input("Введите имя файла №4, в котором содержится ШАБЛОН С КЛЮЧАМИ (с расширением .doc или .docx):\n")
        if file_name4 != "":
            path_file4 = os.path.abspath(file_name4)
            Doc4 = shutil.copy(file_name4, folder_save)

        else:
            path_file4 = " "

        path_save = os.path.abspath(folder_save)
        input_csv(path_file1, path_file2, path_file3, path_file4)

def input_csv(path_file1, path_file2, path_file3, path_file4):
    id = input("Введите id\n")
    doc_name = input("Введите название документа:\n")
    base = input("Введите основание:\n")
    practitioner_data2 = [
        (id, doc_name, base, path_file1, path_file2, path_file3, path_file4)
    ]
    with open("practitioner.csv", "a", newline='', encoding='cp1251', errors="ignore") as csvfile:
        practitioner2 = csv.writer(csvfile, delimiter=";")
        practitioner2.writerows(
            practitioner_data2
        )

def replace_word(sheet):
    for type_name in range(2, sheet.max_row + 1):
        table_key = sheet[type_name][3].value
        if table_key in cell.text:
            table_name = sheet[type_name][2].value
            quest = input("Введите " + table_name + ": ")
            cell.text = cell.text.replace(table_key, quest)

def test_append(table, test1):
    df = [['' for i in range(len(tables[0].columns))] for j in range(len(tables[0].rows))]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell.text:
                df[i][j] = cell.text
                if '${' in cell.text:
                    test1.append(df[i][j])

def final_append(final, test):
    for item in test:
        if item not in final:
            final.append(item)

def str_split(final):
    str_a = ''.join(final)
    str = str_a.split('$')
    str.pop(0)
    for i in range(len(str)):
        str[i] = str[i].split('}')[0]
        str[i] = "$" + str[i] + "}"
    return str

def append_to_excel(str, test2, test3, doc1, vid):
    for i in range(len(str)):
        print(str[i])
        name = input("Название: ")
        doc1.append(k)
        test3.append(name)
        type1 = input("Вид(One or Many): ")
        test2.append(type1)
        new = input("Тип: ")
        vid.append(new)
        del1.append("no")

def create_table(doc1, test3, str, test2, vid, del1):
    name2 = input("Имя файла xlsx(Писать с расширением!): ")
    name3 = input("Имя файла xml(Писать с расширением!): ")
    list = pd.DataFrame({'doc': doc1, 'name': test3, 'key': str, 'type': test2, 'vid': vid, 'del': del1})
    writer = pd.ExcelWriter(name2)
    list.to_excel(writer, sheet_name='main')
    writer.save()
    list.to_xml(name3)


print("\nПривет Оператор, для тебя тут готова программа. Сейчас, тебе нужно выбрать цифру, для того, чтобы продолжить.")


while True:
    number = int(input("\n1 - для оператора 1\n2 - для оператора 2\n3 - для оператора 3\n4 - для выхода из проги\n"))

    if number == 1:
        flag_start = 0
        print("Для оператора 1. Следуйте указаниям!\n")
        flag_csv = int(input("Если csv файл уже создан,то введите 1, иначе 0:\n"))
        flag_check(flag_csv)

        num_of_docs = int(input("Введите количество документов:\n"))
        while flag_start != 1:
            path_check(num_of_docs)
            flag_start = int(input("Если хотите завершить полностью программу, введите 1, если хотите загрузить новый документ, то введите 0:\n"))

    elif number == 2:
        path = input("Полный путь к документу: ")
        doc = Document(path)
        tables = doc.tables
        table = tables[0]
        list = []
        final = []
        sp = []
        test1 = []
        test2 = []
        test3 = []
        del1 = []
        doc1 = []
        vid = []
        test_append(table, test1)
        final_append(final, test1)
        str = str_split(final)
        k = input("Введите ID документа из программы 1: ")

        append_to_excel(str, test2, test3, doc1, vid)

        create_table(doc1, test3, str, test2, vid, del1)

    elif number == 3:
        path_1 = input("Введите путь к файлу .docx: ")
        path_2 = input("Введите путь к файлу .xlsx: ")

        wordDoc = Document(path_1)
        book = openpyxl.open(path_2, read_only=True)
        sheet1 = book.active

        tables1 = wordDoc.tables
        table1 = tables1[0]

        for i, row in enumerate(table1.rows):
            for j, cell in enumerate(row.cells):
                if cell.text:
                    if '${' in cell.text:
                        replace_word(sheet1)

        save_path = input("Введите путь для сохранения: ")
        wordDoc.save(save_path)

    elif number == 4:
        break

    else:
        print("Технические шоколадки")







